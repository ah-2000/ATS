"""
Template Handler Service
Manages resume templates and generates DOCX output
"""

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from typing import Dict, Any, Optional
import re


def create_styled_docx(reconstructed_text: str, filename: str = "reconstructed_resume.docx") -> BytesIO:
    """
    Create a professionally styled DOCX document from reconstructed resume text.
    
    Args:
        reconstructed_text: The reconstructed resume content
        filename: Output filename
        
    Returns:
        BytesIO buffer containing the DOCX file
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Parse and add content
    lines = reconstructed_text.strip().split('\n')
    
    current_section = None
    is_first_line = True
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Detect section headers (=== SECTION === or ALL CAPS with specific keywords)
        if line.startswith('===') and line.endswith('==='):
            # Section header
            section_name = line.replace('=', '').strip()
            add_section_header(doc, section_name)
            current_section = section_name
        elif is_section_header(line):
            add_section_header(doc, line)
            current_section = line
        elif is_first_line or is_name_line(line, is_first_line):
            # Name (first line, typically)
            add_name(doc, line)
            is_first_line = False
        elif is_contact_line(line):
            # Contact info line
            add_contact_line(doc, line)
        elif line.startswith('•') or line.startswith('-') or line.startswith('*'):
            # Bullet point
            add_bullet_point(doc, line[1:].strip())
        elif '|' in line and current_section in ['PROFESSIONAL EXPERIENCE', 'EXPERIENCE', 'WORK EXPERIENCE']:
            # Job title line with company
            add_job_title_line(doc, line)
        elif '|' in line and current_section in ['EDUCATION']:
            # Education line
            add_education_line(doc, line)
        else:
            # Regular paragraph
            add_paragraph(doc, line)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer


def is_section_header(line: str) -> bool:
    """Check if line is a section header"""
    section_keywords = [
        'PROFESSIONAL SUMMARY', 'SUMMARY', 'OBJECTIVE',
        'SKILLS', 'TECHNICAL SKILLS', 'CORE COMPETENCIES',
        'EXPERIENCE', 'PROFESSIONAL EXPERIENCE', 'WORK EXPERIENCE',
        'PROJECTS', 'KEY PROJECTS',
        'EDUCATION', 'ACADEMIC BACKGROUND',
        'CERTIFICATIONS', 'CERTIFICATES',
        'LANGUAGES', 'ADDITIONAL INFORMATION',
        'HEADER'
    ]
    cleaned = line.replace('=', '').strip().upper()
    return cleaned in section_keywords


def is_name_line(line: str, is_first: bool) -> bool:
    """Check if line is likely a name"""
    if not is_first:
        return False
    # Names are typically short, no special chars except spaces
    if len(line) < 50 and not any(c in line for c in ['@', '|', '•', '-', ':', '/']):
        return True
    return False


def is_contact_line(line: str) -> bool:
    """Check if line contains contact information"""
    indicators = ['@', '|', 'linkedin', 'github', '+1', '+92', 'phone', 'email']
    line_lower = line.lower()
    return any(ind in line_lower for ind in indicators)


def add_name(doc: Document, name: str):
    """Add name as document title"""
    para = doc.add_paragraph()
    run = para.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x2E, 0x2E, 0x2E)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(2)


def add_contact_line(doc: Document, contact: str):
    """Add contact information line"""
    para = doc.add_paragraph()
    run = para.add_run(contact)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.space_after = Pt(12)


def add_section_header(doc: Document, header: str):
    """Add section header with styling"""
    # Clean header
    header = header.replace('=', '').strip().upper()
    if header == 'HEADER':
        return  # Skip the HEADER section marker
    
    para = doc.add_paragraph()
    para.space_before = Pt(12)
    run = para.add_run(header)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x6C)  # Dark blue
    
    # Add underline effect via border
    para.paragraph_format.space_after = Pt(4)


def add_job_title_line(doc: Document, line: str):
    """Add job title and company line"""
    para = doc.add_paragraph()
    para.space_before = Pt(8)
    
    parts = line.split('|')
    for i, part in enumerate(parts):
        part = part.strip()
        run = para.add_run(part)
        if i == 0:  # Job title - bold
            run.bold = True
            run.font.size = Pt(11)
        else:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        
        if i < len(parts) - 1:
            para.add_run(' | ')
    
    para.space_after = Pt(2)


def add_education_line(doc: Document, line: str):
    """Add education line"""
    para = doc.add_paragraph()
    para.space_before = Pt(4)
    
    parts = line.split('|')
    for i, part in enumerate(parts):
        part = part.strip()
        run = para.add_run(part)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True
        
        if i < len(parts) - 1:
            para.add_run(' | ')
    
    para.space_after = Pt(2)


def add_bullet_point(doc: Document, text: str):
    """Add bullet point"""
    para = doc.add_paragraph(style='List Bullet')
    run = para.add_run(text)
    run.font.size = Pt(10)
    para.paragraph_format.left_indent = Inches(0.25)
    para.space_after = Pt(2)


def add_paragraph(doc: Document, text: str):
    """Add regular paragraph"""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(10)
    para.space_after = Pt(4)


def get_default_template_info() -> Dict[str, Any]:
    """Get information about the default template"""
    return {
        "name": "Professional Clean Template",
        "description": "Clean and professional resume template with clear sections",
        "sections": [
            "Header (Name & Contact)",
            "Professional Summary",
            "Skills",
            "Professional Experience",
            "Projects",
            "Education",
            "Certifications",
            "Languages"
        ],
        "formatting": {
            "margins": "0.75 inch",
            "name_size": "18pt bold",
            "section_headers": "12pt bold dark blue",
            "body_text": "10pt regular"
        }
    }

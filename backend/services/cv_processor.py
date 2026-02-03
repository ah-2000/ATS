"""
CV Processor Service
Handles CV parsing, analysis, enhancement, and regeneration with formatting preservation.
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
import json
import pdfplumber
import copy


# ============================================
# TEXT EXTRACTION
# ============================================

def extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from PDF file."""
    text = ""
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
    except Exception as e:
        raise Exception(f"Error reading PDF: {str(e)}")
    return text


def extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        doc = Document(BytesIO(file_bytes))
        text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                text += "\n" + row_text
        
        return text
    except Exception as e:
        raise Exception(f"Error reading DOCX: {str(e)}")


def extract_text(file_bytes: bytes, file_type: str) -> str:
    """Extract text from file based on type."""
    if file_type == "pdf":
        return extract_pdf_text(file_bytes)
    elif file_type == "docx":
        return extract_docx_text(file_bytes)
    else:
        raise Exception(f"Unsupported file type: {file_type}")


# ============================================
# CV ANALYSIS
# ============================================

def get_analysis_prompt(cv_text: str, job_description: str, job_position: str) -> str:
    """Generate the ATS analysis prompt."""
    return f'''Hey, act like a skilled and experienced ATS (Application Tracking System) with deep understanding of both technical and non-technical fields.
Your task is to evaluate the resume based on the given job description for the job position "{job_position}".
You must consider the job market is very competitive and provide the best assistance for improving resumes.
Assign the percentage matching based on the job description and list the missing keywords with high accuracy.
Evaluate the resume while ignoring name, gender, and age.

**Requirements:**
- Assign a JD match percentage (accurate score).
- Highlight missing skills (only relevant ones).
- Extract Key Strengths and provide Recommendations.
- Ensure the response **always contains** a "Profile Summary".
- Provide weighted scoring breakdown.

**Input:**
Job Position: {job_position}
Resume: {cv_text}
Job Description: {job_description}

**Return Response in Strict JSON Format:**
{{
    "JD Match": "XX%",
    "MissingKeywords": ["Skill1", "Skill2", "Skill3"],
    "KeyStrength": "Brief summary of key strengths",
    "Recommendations": "Brief recommendations for improvement",
    "Profile Summary": "Concise evaluation of strengths and gaps.",
    "ExperienceMatch": "XX%",
    "SkillsMatch": "XX%",
    "EducationMatch": "XX%"
}}'''


def parse_json_response(response: str) -> Optional[Dict[str, Any]]:
    """Parse and clean AI response to extract JSON."""
    try:
        response = response.strip()
        
        # Remove markdown code blocks
        if response.startswith("```json"):
            response = response[7:]
        elif response.startswith("```"):
            response = response[3:]
        
        if response.endswith("```"):
            response = response[:-3]
        
        response = response.strip()
        
        # Try to find JSON object if response contains extra text
        if not response.startswith('{'):
            start_idx = response.find('{')
            if start_idx != -1:
                response = response[start_idx:]
        
        if not response.endswith('}'):
            end_idx = response.rfind('}')
            if end_idx != -1:
                response = response[:end_idx + 1]
        
        parsed = json.loads(response)
        print(f"Successfully parsed JSON response: {list(parsed.keys())}")
        return parsed
        
    except json.JSONDecodeError as e:
        print(f"JSON Parse Error: {str(e)}")
        print(f"Response preview (first 500 chars): {response[:500]}")
        return None
    except Exception as e:
        print(f"Unexpected error in parse_json_response: {str(e)}")
        return None


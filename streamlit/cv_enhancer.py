"""
CV Enhancement Module
Handles parsing, AI enhancement, and regeneration of CVs with formatting preservation.
"""

import fitz  # PyMuPDF
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, List, Any, Optional, Tuple
import json
import re
import copy


# ============================================
# DOCX PARSING AND ENHANCEMENT
# ============================================

def extract_docx_content_with_formatting(file_bytes: BytesIO) -> Tuple[Document, List[Dict[str, Any]]]:
    """
    Extract content from DOCX with formatting metadata.
    
    Returns:
        Tuple of (Document object, list of content sections with formatting info)
    """
    doc = Document(file_bytes)
    sections = []
    
    for para_idx, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            para_info = {
                "type": "paragraph",
                "index": para_idx,
                "text": paragraph.text,
                "style": paragraph.style.name if paragraph.style else "Normal",
                "alignment": str(paragraph.alignment) if paragraph.alignment else "LEFT",
                "runs": []
            }
            
            for run_idx, run in enumerate(paragraph.runs):
                if run.text:
                    run_info = {
                        "index": run_idx,
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                        "font_name": run.font.name,
                        "font_size": str(run.font.size) if run.font.size else None,
                    }
                    para_info["runs"].append(run_info)
            
            sections.append(para_info)
    
    # Also extract tables
    for table_idx, table in enumerate(doc.tables):
        table_info = {
            "type": "table",
            "index": table_idx,
            "rows": []
        }
        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                row_cells.append(cell.text)
            table_info["rows"].append(row_cells)
        sections.append(table_info)
    
    return doc, sections


def get_docx_text_content(sections: List[Dict[str, Any]]) -> str:
    """Extract plain text from sections for AI analysis."""
    text_parts = []
    for section in sections:
        if section["type"] == "paragraph":
            text_parts.append(section["text"])
        elif section["type"] == "table":
            for row in section["rows"]:
                text_parts.append(" | ".join(row))
    return "\n".join(text_parts)


def apply_enhancements_to_docx(doc: Document, enhancements: List[Dict[str, str]]) -> Document:
    """
    Apply AI enhancements to DOCX while preserving formatting.
    
    Args:
        doc: Original Document object
        enhancements: List of {original: str, improved: str} mappings
    
    Returns:
        Modified Document object
    """
    # Create a mapping for quick lookup
    enhancement_map = {e["original"].strip(): e["improved"] for e in enhancements}
    
    # Process paragraphs
    for paragraph in doc.paragraphs:
        para_text = paragraph.text.strip()
        
        # Check if this paragraph needs enhancement
        if para_text in enhancement_map:
            improved_text = enhancement_map[para_text]
            
            # Strategy: Replace text in the first run, clear others
            if paragraph.runs:
                # Keep formatting from first run
                first_run = paragraph.runs[0]
                first_run.text = improved_text
                
                # Clear remaining runs
                for run in paragraph.runs[1:]:
                    run.text = ""
        else:
            # Check for partial matches within runs
            for original, improved in enhancement_map.items():
                if original in para_text:
                    for run in paragraph.runs:
                        if original in run.text:
                            run.text = run.text.replace(original, improved)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text in enhancement_map:
                    # Preserve formatting by working with paragraphs in cell
                    for para in cell.paragraphs:
                        if para.runs:
                            para.runs[0].text = enhancement_map[cell_text]
                            for run in para.runs[1:]:
                                run.text = ""
    
    return doc


def generate_improved_docx(original_bytes: BytesIO, enhancements: List[Dict[str, str]]) -> BytesIO:
    """
    Generate improved DOCX file with enhancements applied.
    
    Args:
        original_bytes: Original DOCX file as BytesIO
        enhancements: List of enhancement mappings from AI
    
    Returns:
        BytesIO containing the improved DOCX
    """
    # Reset file pointer
    original_bytes.seek(0)
    
    # Load and enhance
    doc = Document(original_bytes)
    enhanced_doc = apply_enhancements_to_docx(doc, enhancements)
    
    # Save to BytesIO
    output = BytesIO()
    enhanced_doc.save(output)
    output.seek(0)
    
    return output


# ============================================
# PDF PARSING AND ENHANCEMENT
# ============================================

def extract_pdf_content_with_formatting(file_bytes: BytesIO) -> Tuple[fitz.Document, List[Dict[str, Any]]]:
    """
    Extract content from PDF with formatting metadata.
    
    Returns:
        Tuple of (PDF document, list of text blocks with formatting info)
    """
    file_bytes.seek(0)
    pdf_doc = fitz.open(stream=file_bytes.read(), filetype="pdf")
    
    sections = []
    
    for page_num, page in enumerate(pdf_doc):
        # Get text blocks with detailed info
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        
        for block_idx, block in enumerate(blocks):
            if block["type"] == 0:  # Text block
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        section = {
                            "type": "text",
                            "page": page_num,
                            "text": span.get("text", ""),
                            "font": span.get("font", ""),
                            "size": span.get("size", 12),
                            "color": span.get("color", 0),
                            "bbox": span.get("bbox", [0, 0, 0, 0]),
                            "flags": span.get("flags", 0),  # bold, italic, etc.
                        }
                        if section["text"].strip():
                            sections.append(section)
    
    return pdf_doc, sections


def get_pdf_text_content(sections: List[Dict[str, Any]]) -> str:
    """Extract plain text from PDF sections for AI analysis."""
    # Group by page and sort by position
    pages = {}
    for section in sections:
        page_num = section.get("page", 0)
        if page_num not in pages:
            pages[page_num] = []
        pages[page_num].append(section)
    
    text_parts = []
    for page_num in sorted(pages.keys()):
        page_sections = sorted(pages[page_num], key=lambda x: (x["bbox"][1], x["bbox"][0]))
        for section in page_sections:
            text_parts.append(section["text"])
    
    return "\n".join(text_parts)


def apply_enhancements_to_pdf(pdf_bytes: BytesIO, enhancements: List[Dict[str, str]]) -> BytesIO:
    """
    Apply AI enhancements to PDF.
    
    Note: PDF editing while preserving exact formatting is challenging.
    This uses redaction + text insertion approach.
    
    Args:
        pdf_bytes: Original PDF as BytesIO
        enhancements: List of {original: str, improved: str} mappings
    
    Returns:
        BytesIO containing the improved PDF
    """
    pdf_bytes.seek(0)
    pdf_doc = fitz.open(stream=pdf_bytes.read(), filetype="pdf")
    
    enhancement_map = {e["original"].strip(): e["improved"] for e in enhancements}
    
    for page in pdf_doc:
        for original, improved in enhancement_map.items():
            # Find all instances of original text
            text_instances = page.search_for(original)
            
            for rect in text_instances:
                # Get font info from the area (approximate)
                # Extract text details to get font size
                text_dict = page.get_text("dict", clip=rect)
                font_size = 11  # default
                font_name = "helv"  # default (Helvetica)
                
                for block in text_dict.get("blocks", []):
                    if block["type"] == 0:
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                font_size = span.get("size", 11)
                                break
                
                # Add redaction annotation
                page.add_redact_annot(rect, fill=(1, 1, 1))  # White fill
                
        # Apply all redactions at once
        page.apply_redactions()
        
        # Now insert the new text
        for original, improved in enhancement_map.items():
            text_instances = page.search_for(original)
            # Note: After redaction, original text is gone, so we need to 
            # remember positions before redacting
    
    # Alternative approach: Create a new PDF with text replacement
    # This is more reliable for complex cases
    pdf_bytes.seek(0)
    pdf_doc = fitz.open(stream=pdf_bytes.read(), filetype="pdf")
    
    for page_num in range(len(pdf_doc)):
        page = pdf_doc[page_num]
        
        for original, improved in enhancement_map.items():
            # Find text instances
            instances = page.search_for(original)
            
            if instances:
                for inst in instances:
                    # Get the text properties
                    blocks = page.get_text("dict", clip=inst)["blocks"]
                    font_size = 11
                    font_color = (0, 0, 0)
                    
                    for block in blocks:
                        if block["type"] == 0:
                            for line in block.get("lines", []):
                                for span in line.get("spans", []):
                                    font_size = span.get("size", 11)
                                    color_int = span.get("color", 0)
                                    # Convert integer color to RGB
                                    r = ((color_int >> 16) & 255) / 255
                                    g = ((color_int >> 8) & 255) / 255
                                    b = (color_int & 255) / 255
                                    font_color = (r, g, b)
                                    break
                    
                    # Redact the original text
                    page.add_redact_annot(inst, fill=(1, 1, 1))
                    page.apply_redactions()
                    
                    # Insert improved text at the same position
                    # Adjust position slightly for better alignment
                    text_point = fitz.Point(inst.x0, inst.y1 - 2)
                    page.insert_text(
                        text_point,
                        improved,
                        fontsize=font_size,
                        color=font_color
                    )
    
    # Save to BytesIO
    output = BytesIO()
    pdf_doc.save(output)
    output.seek(0)
    
    return output


def generate_improved_pdf(original_bytes: BytesIO, enhancements: List[Dict[str, str]]) -> BytesIO:
    """
    Generate improved PDF file with enhancements applied.
    
    Args:
        original_bytes: Original PDF file as BytesIO
        enhancements: List of enhancement mappings from AI
    
    Returns:
        BytesIO containing the improved PDF
    """
    return apply_enhancements_to_pdf(original_bytes, enhancements)


# ============================================
# AI ENHANCEMENT PROMPT
# ============================================

def get_cv_enhancement_prompt(cv_text: str, job_description: str, job_position: str, analysis_result: Dict[str, Any]) -> str:
    """
    Generate the AI prompt for CV content enhancement.
    
    Args:
        cv_text: Plain text content of the CV
        job_description: Target job description
        job_position: Target job position
        analysis_result: Previous ATS analysis result containing recommendations
    
    Returns:
        Formatted prompt string
    """
    prompt = f'''You are an expert CV/Resume writer. Your task is to improve the CV content to better match the job requirements while maintaining the candidate's authentic experience.

**Job Position:** {job_position}

**Job Description:**
{job_description}

**Previous ATS Analysis:**
- JD Match: {analysis_result.get("JD Match", "N/A")}
- Missing Keywords: {", ".join(analysis_result.get("MissingKeywords", []))}
- Recommendations: {analysis_result.get("Recommendations", "N/A")}

**Original CV Content:**
{cv_text}

**Your Task:**
1. Identify bullet points, descriptions, and sections that can be improved
2. Rewrite weak content with stronger, impact-focused language
3. Incorporate relevant missing keywords naturally where appropriate
4. Quantify achievements where possible (add realistic metrics if not present)
5. Improve clarity and professional tone
6. DO NOT fabricate experience or skills the candidate doesn't have
7. Keep section headers and structure unchanged

**IMPORTANT RULES:**
- Only suggest changes for content that genuinely needs improvement
- Preserve the exact wording of section headers (Education, Experience, Skills, etc.)
- Match the original text EXACTLY when specifying what to replace
- Each "original" field must contain the EXACT text from the CV

**Return your response in this exact JSON format:**
{{
    "enhancements": [
        {{
            "section": "Section name (e.g., Experience, Skills)",
            "original": "The exact original text to replace",
            "improved": "The improved version of the text",
            "reason": "Brief reason for this improvement"
        }}
    ],
    "summary": "Brief summary of all improvements made"
}}

Return ONLY valid JSON, no markdown formatting or additional text.'''
    
    return prompt


def parse_enhancement_response(response: str) -> Optional[Dict[str, Any]]:
    """
    Parse the AI enhancement response.
    
    Args:
        response: Raw AI response string
    
    Returns:
        Parsed enhancement data or None if parsing fails
    """
    try:
        # Clean up response
        response = response.strip()
        
        # Remove markdown code blocks if present
        if response.startswith("```json"):
            response = response[7:]
        if response.startswith("```"):
            response = response[3:]
        if response.endswith("```"):
            response = response[:-3]
        
        response = response.strip()
        
        return json.loads(response)
    except json.JSONDecodeError:
        return None


# ============================================
# MAIN ENHANCEMENT FUNCTION
# ============================================

def enhance_cv(
    file_bytes: BytesIO,
    file_type: str,
    job_description: str,
    job_position: str,
    analysis_result: Dict[str, Any],
    ai_response_func,
    provider: str,
    model: str
) -> Tuple[Optional[BytesIO], str, Optional[Dict[str, Any]]]:
    """
    Main function to enhance a CV.
    
    Args:
        file_bytes: Original file as BytesIO
        file_type: "pdf" or "docx"
        job_description: Target job description
        job_position: Target job position
        analysis_result: Previous ATS analysis result
        ai_response_func: Function to call AI (get_ai_response)
        provider: AI provider name
        model: AI model name
    
    Returns:
        Tuple of (enhanced_file_bytes, status_message, enhancement_data)
    """
    try:
        # Extract content based on file type
        file_bytes.seek(0)
        
        if file_type == "docx":
            doc, sections = extract_docx_content_with_formatting(file_bytes)
            cv_text = get_docx_text_content(sections)
        elif file_type == "pdf":
            pdf_doc, sections = extract_pdf_content_with_formatting(file_bytes)
            cv_text = get_pdf_text_content(sections)
            pdf_doc.close()
        else:
            return None, f"Unsupported file type: {file_type}", None
        
        # Generate enhancement prompt
        prompt = get_cv_enhancement_prompt(cv_text, job_description, job_position, analysis_result)
        
        # Get AI response
        ai_response = ai_response_func(prompt, provider, model)
        
        # Parse response
        enhancement_data = parse_enhancement_response(ai_response)
        
        if not enhancement_data or "enhancements" not in enhancement_data:
            return None, "Failed to parse AI enhancement suggestions", None
        
        enhancements = enhancement_data["enhancements"]
        
        if not enhancements:
            return None, "No improvements suggested - your CV is already well-optimized!", enhancement_data
        
        # Apply enhancements
        file_bytes.seek(0)
        
        if file_type == "docx":
            enhanced_bytes = generate_improved_docx(file_bytes, enhancements)
        else:
            enhanced_bytes = generate_improved_pdf(file_bytes, enhancements)
        
        summary = enhancement_data.get("summary", f"Applied {len(enhancements)} improvements")
        
        return enhanced_bytes, summary, enhancement_data
        
    except Exception as e:
        return None, f"Error enhancing CV: {str(e)}", None
